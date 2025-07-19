import os
import json
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches
import weasyprint
from jinja2 import Template

class ReportGenerator:
    """Service to generate final equity research reports"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def generate_report(self, company_name, quarter, financial_data, transcript_analysis):
        """
        Generate complete equity research report
        Returns HTML report content
        """
        try:
            # Prepare report data
            report_data = {
                'company_name': company_name,
                'quarter': quarter,
                'generated_date': datetime.now().strftime('%B %d, %Y'),
                'financial_data': financial_data,
                'transcript_analysis': transcript_analysis,
                'executive_summary': self._generate_executive_summary(
                    company_name, quarter, financial_data, transcript_analysis
                ),
                'financial_summary': self._create_financial_summary_table(financial_data),
                'key_metrics': self._extract_key_metrics(financial_data),
                'investment_thesis': self._generate_investment_thesis(transcript_analysis)
            }
            
            # Generate HTML report
            html_report = self._generate_html_report(report_data)
            
            self.logger.info(f"Successfully generated report for {company_name} {quarter}")
            return html_report
            
        except Exception as e:
            self.logger.error(f"Error generating report: {str(e)}")
            raise Exception(f"Failed to generate report: {str(e)}")
    
    def export_to_pdf(self, html_content, company_name, quarter, download_folder):
        """Export report to PDF format"""
        try:
            filename = f"{company_name}_{quarter}_Research_Report_{datetime.now().strftime('%Y%m%d')}.pdf"
            file_path = os.path.join(download_folder, filename)
            
            # Convert HTML to PDF using weasyprint
            weasyprint.HTML(string=html_content).write_pdf(file_path)
            
            self.logger.info(f"Successfully exported PDF report: {file_path}")
            return file_path
            
        except Exception as e:
            self.logger.error(f"Error exporting to PDF: {str(e)}")
            raise Exception(f"Failed to export PDF: {str(e)}")
    
    def export_to_docx(self, html_content, company_name, quarter, download_folder):
        """Export report to Word document format"""
        try:
            filename = f"{company_name}_{quarter}_Research_Report_{datetime.now().strftime('%Y%m%d')}.docx"
            file_path = os.path.join(download_folder, filename)
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'{company_name} - {quarter} Equity Research Report', 0)
            
            # Add generation date
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
            doc.add_paragraph('')
            
            # Parse HTML content and add to document
            # This is a simplified conversion - in production, you might want a more sophisticated HTML to DOCX converter
            self._add_html_content_to_docx(doc, html_content)
            
            # Save document
            doc.save(file_path)
            
            self.logger.info(f"Successfully exported Word document: {file_path}")
            return file_path
            
        except Exception as e:
            self.logger.error(f"Error exporting to Word: {str(e)}")
            raise Exception(f"Failed to export Word document: {str(e)}")
    
    def _generate_executive_summary(self, company_name, quarter, financial_data, transcript_analysis):
        """Generate executive summary"""
        try:
            # Extract key financial metrics
            key_metrics = self._extract_key_metrics(financial_data)
            
            summary_parts = []
            
            # Financial performance summary
            if key_metrics:
                summary_parts.append(f"{company_name} reported {quarter} results with key metrics including revenue and profitability measures. The financial performance shows the company's operational execution during the quarter.")
            
            # Strategic themes from transcript
            if transcript_analysis.get('strategic_themes'):
                summary_parts.append(f"Management highlighted strategic initiatives including {transcript_analysis['strategic_themes'][:200]}...")
            
            # Risk and outlook
            if transcript_analysis.get('risks_and_tailwinds'):
                summary_parts.append(f"Key considerations for investors include {transcript_analysis['risks_and_tailwinds'][:200]}...")
            
            return ' '.join(summary_parts) if summary_parts else f"Analysis of {company_name}'s {quarter} earnings results and management commentary."
            
        except Exception as e:
            self.logger.warning(f"Error generating executive summary: {str(e)}")
            return f"Executive summary for {company_name} {quarter} earnings analysis."
    
    def _create_financial_summary_table(self, financial_data):
        """Create formatted financial summary table"""
        if not financial_data or 'line_items' not in financial_data:
            return []
        
        table_data = []
        for item in financial_data['line_items']:
            row = {
                'metric': item['line_item'],
                'previous': f"{item['previous_value']:,.0f}" if item.get('previous_value') else 'N/A',
                'estimate': f"{item['estimate']:,.0f}" if item.get('estimate') else 'N/A',
                'actual': 'TBD',  # Placeholder for actual values
                'variance': 'TBD'  # Placeholder for variance calculation
            }
            table_data.append(row)
        
        return table_data
    
    def _extract_key_metrics(self, financial_data):
        """Extract key financial metrics"""
        if not financial_data or 'line_items' not in financial_data:
            return []
        
        # Return first 5 line items as key metrics
        return financial_data['line_items'][:5]
    
    def _generate_investment_thesis(self, transcript_analysis):
        """Generate investment thesis based on transcript analysis"""
        thesis_points = []
        
        if transcript_analysis.get('strategic_themes'):
            thesis_points.append(f"Strategic Focus: {transcript_analysis['strategic_themes']}")
        
        if transcript_analysis.get('market_dynamics'):
            thesis_points.append(f"Market Position: {transcript_analysis['market_dynamics']}")
        
        if transcript_analysis.get('financial_highlights'):
            thesis_points.append(f"Financial Strength: {transcript_analysis['financial_highlights']}")
        
        return thesis_points
    
    def _generate_html_report(self, report_data):
        """Generate HTML report using template"""
        template_str = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{{ company_name }} - {{ quarter }} Research Report</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; color: #333; }
                .header { border-bottom: 2px solid #007bff; padding-bottom: 20px; margin-bottom: 30px; }
                .company-name { font-size: 28px; font-weight: bold; color: #007bff; }
                .quarter { font-size: 20px; color: #666; }
                .section { margin-bottom: 30px; }
                .section-title { font-size: 18px; font-weight: bold; color: #007bff; margin-bottom: 15px; border-bottom: 1px solid #ddd; padding-bottom: 5px; }
                .financial-table { width: 100%; border-collapse: collapse; margin-bottom: 20px; }
                .financial-table th, .financial-table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                .financial-table th { background-color: #f8f9fa; font-weight: bold; }
                .content { line-height: 1.6; }
                .metric-item { margin-bottom: 10px; }
                .metric-label { font-weight: bold; }
                .footer { margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; color: #666; font-size: 12px; }
            </style>
        </head>
        <body>
            <div class="header">
                <div class="company-name">{{ company_name }}</div>
                <div class="quarter">{{ quarter }} Equity Research Report</div>
                <div style="margin-top: 10px; color: #666;">Generated on {{ generated_date }}</div>
            </div>

            <div class="section">
                <div class="section-title">Executive Summary</div>
                <div class="content">{{ executive_summary }}</div>
            </div>

            <div class="section">
                <div class="section-title">Financial Summary</div>
                {% if financial_summary %}
                <table class="financial-table">
                    <thead>
                        <tr>
                            <th>Metric</th>
                            <th>Previous Quarter</th>
                            <th>Analyst Estimate</th>
                            <th>Actual</th>
                            <th>Variance</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for row in financial_summary %}
                        <tr>
                            <td>{{ row.metric }}</td>
                            <td>{{ row.previous }}</td>
                            <td>{{ row.estimate }}</td>
                            <td>{{ row.actual }}</td>
                            <td>{{ row.variance }}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
                {% else %}
                <p>Financial data not available.</p>
                {% endif %}
            </div>

            <div class="section">
                <div class="section-title">Management Commentary</div>
                <div class="content">{{ transcript_analysis.management_commentary or 'No commentary available.' }}</div>
            </div>

            <div class="section">
                <div class="section-title">Strategic Themes</div>
                <div class="content">{{ transcript_analysis.strategic_themes or 'No strategic themes identified.' }}</div>
            </div>

            <div class="section">
                <div class="section-title">Risks and Tailwinds</div>
                <div class="content">{{ transcript_analysis.risks_and_tailwinds or 'No risks or tailwinds identified.' }}</div>
            </div>

            <div class="section">
                <div class="section-title">Q&A Insights</div>
                <div class="content">{{ transcript_analysis.qa_insights or 'No Q&A insights available.' }}</div>
            </div>

            {% if investment_thesis %}
            <div class="section">
                <div class="section-title">Investment Thesis</div>
                <div class="content">
                    {% for point in investment_thesis %}
                    <div class="metric-item">• {{ point }}</div>
                    {% endfor %}
                </div>
            </div>
            {% endif %}

            <div class="footer">
                <p><strong>Disclaimer:</strong> This research report is generated using automated analysis tools and should be used for informational purposes only. Please conduct your own due diligence before making investment decisions.</p>
                <p>Report generated on {{ generated_date }} using AI-powered equity research tools.</p>
            </div>
        </body>
        </html>
        """
        
        template = Template(template_str)
        return template.render(**report_data)
    
    def _add_html_content_to_docx(self, doc, html_content):
        """Add simplified HTML content to Word document"""
        try:
            # This is a basic implementation - in production you might want to use a proper HTML to DOCX converter
            from bs4 import BeautifulSoup
            
            # Parse HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract sections and add to document
            sections = soup.find_all(class_='section')
            for section in sections:
                title = section.find(class_='section-title')
                content = section.find(class_='content')
                
                if title:
                    doc.add_heading(title.get_text(), level=1)
                
                if content:
                    doc.add_paragraph(content.get_text())
                
                doc.add_paragraph('')  # Add spacing
                
        except Exception as e:
            self.logger.warning(f"Error parsing HTML for Word document: {str(e)}")
            # Fallback: add plain text version
            doc.add_paragraph("Report content could not be formatted properly. Please refer to the PDF version.")
